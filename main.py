import os
import uuid
import base64
import sqlite3
import hashlib
import requests
from io import BytesIO
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document

# ========================
# CONFIGURATION
# ========================
GEMINI_API_KEY = "AIzaSyCdK2hdmiF4m4XS3IeiR3WDWcmg6WoRgWM"
GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = FastAPI(title="RAG + Analyse Document IA")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ========================
# BASE DE DONNÉES SIMPLE
# ========================
def init_db():
    conn = sqlite3.connect("users.db")
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def register_user(username, password):
    try:
        conn = sqlite3.connect("users.db")
        cursor = conn.cursor()
        hashed_password = hashlib.sha256(password.encode()).hexdigest()
        cursor.execute('INSERT INTO users (username, password) VALUES (?, ?)', (username, hashed_password))
        conn.commit()
        conn.close()
        return True
    except:
        return False

def login_user(username, password):
    conn = sqlite3.connect("users.db")
    cursor = conn.cursor()
    hashed_password = hashlib.sha256(password.encode()).hexdigest()
    cursor.execute('SELECT id FROM users WHERE username = ? AND password = ?', (username, hashed_password))
    user = cursor.fetchone()
    conn.close()
    return user is not None

# Initialiser la base
init_db()

# Stockage simple des sessions
sessions = {}
user_data = {}  # Stocke les données par utilisateur

# ========================
# UTILITAIRES (identique à ton code perfection)
# ========================
def extract_text_from_file(path: str, ext: str) -> str:
    ext = ext.lower()
    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == ".csv":
        return pd.read_csv(path).to_string()
    if ext == ".docx":
        return "\n".join([p.text for p in Document(path).paragraphs])
    if ext in (".xls", ".xlsx"):
        return pd.read_excel(path).to_string()
    raise ValueError("Format non supporté")

def summarize_with_gemini(text: str, max_chars: int = 8000) -> str:
    headers = {"Content-Type": "application/json"}
    prompt = f"Résume ce contenu en français, en 5–8 phrases simples et actionnables :\n\n{text[:max_chars]}"
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    try:
        resp = requests.post(GEMINI_URL, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        return f"Résumé indisponible ({e})."

def gemini_chat(messages: list) -> str:
    headers = {"Content-Type": "application/json"}
    contents = [{"parts": [{"text": msg["content"]}]} for msg in messages]
    payload = {"contents": contents}
    try:
        resp = requests.post(GEMINI_URL, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        return data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
    except Exception as e:
        return f"Impossible d'obtenir une réponse ({e})."

def detect_excel_errors(df: pd.DataFrame) -> str:
    if df is None:
        return ""
    empty = int(df.isna().sum().sum())
    dup = int(df.duplicated().sum())
    msg = []
    if empty:
        msg.append(f"{empty} cellule(s) vide(s)")
    if dup:
        msg.append(f"{dup} ligne(s) en double")
    for col in df.columns:
        if df[col].isna().any():
            msg.append(f"Colonne '{col}' : {df[col].isna().sum()} valeur(s) manquante(s)")
        if df[col].duplicated().any():
            msg.append(f"Colonne '{col}' : {df[col].duplicated().sum()} doublon(s)")
    return " ; ".join(msg) if msg else "Aucune anomalie détectée."

def correct_excel(df: pd.DataFrame, save_path: str) -> str:
    df_corrected = df.copy()
    for col in df_corrected.select_dtypes(include=["object"]).columns:
        try:
            series = df_corrected[col].astype(str).str.strip()
            series = series.str.replace(" ", "", regex=False).str.replace(",", ".", regex=False)
            numeric = pd.to_numeric(series, errors="coerce")
            if numeric.notna().sum() > 0 and numeric.notna().sum() >= (0.2 * len(series)):
                df_corrected[col] = numeric
        except Exception:
            pass
    for col in df_corrected.select_dtypes(include=["float64", "int64"]).columns:
        if df_corrected[col].isna().any():
            df_corrected[col] = df_corrected[col].fillna(df_corrected[col].mean())
    df_corrected.drop_duplicates(inplace=True)
    if not save_path.lower().endswith(".xlsx"):
        save_path += ".xlsx"
    df_corrected.to_excel(save_path, index=False)
    return save_path

def correct_docx(path: str, save_path: str) -> str:
    doc = Document(path)
    new_doc = Document()
    for para in doc.paragraphs:
        text = para.text
        try:
            corrected_text = summarize_with_gemini("Corrige ce texte sans changer la structure :\n" + text)
        except Exception:
            corrected_text = text
        new_doc.add_paragraph(corrected_text)
    if not save_path.lower().endswith(".docx"):
        save_path += ".docx"
    new_doc.save(save_path)
    return save_path

def dataframe_chart_base64(df: pd.DataFrame) -> Optional[str]:
    col = next((c for c in ("Statut", "Status", "state") if c in df.columns), None)
    if not col:
        return None
    counts = df[col].value_counts()
    plt.figure(figsize=(6, 3.5))
    counts.plot(kind="bar")
    plt.title("Répartition par statut")
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=120)
    plt.close()
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")

def get_username_from_session(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id in sessions:
        return sessions[session_id]
    return None

# ========================
# PAGES AUTHENTIFICATION SIMPLES
# ========================

LOGIN_HTML = """
<!DOCTYPE html>
<html>
<head>
    <title>Connexion</title>
    <style>
        body { font-family: Arial; margin: 40px; background: #f0f0f0; }
        .container { max-width: 400px; margin: 0 auto; background: white; padding: 20px; border-radius: 10px; }
        input, button { width: 100%; padding: 10px; margin: 5px 0; }
        button { background: #007bff; color: white; border: none; cursor: pointer; }
        .error { color: red; }
        .links { text-align: center; margin-top: 10px; }
    </style>
</head>
<body>
    <div class="container">
        <h2>Connexion</h2>
        <form method="post" action="/login">
            <input type="text" name="username" placeholder="Nom d'utilisateur" required>
            <input type="password" name="password" placeholder="Mot de passe" required>
            <button type="submit">Se connecter</button>
        </form>
        <div class="links">
            <a href="/register">Créer un compte</a>
        </div>
        <div class="error">{error}</div>
    </div>
</body>
</html>
"""

REGISTER_HTML = """
<!DOCTYPE html>
<html>
<head>
    <title>Inscription</title>
    <style>
        body { font-family: Arial; margin: 40px; background: #f0f0f0; }
        .container { max-width: 400px; margin: 0 auto; background: white; padding: 20px; border-radius: 10px; }
        input, button { width: 100%; padding: 10px; margin: 5px 0; }
        button { background: #28a745; color: white; border: none; cursor: pointer; }
        .error { color: red; }
        .links { text-align: center; margin-top: 10px; }
    </style>
</head>
<body>
    <div class="container">
        <h2>Inscription</h2>
        <form method="post" action="/register">
            <input type="text" name="username" placeholder="Nom d'utilisateur" required>
            <input type="password" name="password" placeholder="Mot de passe" required>
            <button type="submit">S'inscrire</button>
        </form>
        <div class="links">
            <a href="/">Se connecter</a>
        </div>
        <div class="error">{error}</div>
    </div>
</body>
</html>
"""

# ========================
# ENDPOINTS AUTHENTIFICATION
# ========================

@app.get("/")
async def login_page():
    return HTMLResponse(LOGIN_HTML.replace("{error}", ""))

@app.get("/register")
async def register_page():
    return HTMLResponse(REGISTER_HTML.replace("{error}", ""))

@app.post("/register")
async def register_user(request: Request, username: str = Form(...), password: str = Form(...)):
    if register_user(username, password):
        return RedirectResponse(url="/", status_code=303)
    else:
        return HTMLResponse(REGISTER_HTML.replace("{error}", "Nom d'utilisateur déjà utilisé"))

@app.post("/login")
async def login_user(request: Request, username: str = Form(...), password: str = Form(...)):
    if login_user(username, password):
        session_id = str(uuid.uuid4())
        sessions[session_id] = username
        # Initialiser les données utilisateur
        user_data[username] = {
            "last_uploaded_text": "",
            "last_uploaded_df": None,
            "last_corrected_file": None
        }
        response = RedirectResponse(url="/app", status_code=303)
        response.set_cookie(key="session_id", value=session_id)
        return response
    else:
        return HTMLResponse(LOGIN_HTML.replace("{error}", "Identifiants incorrects"))

@app.get("/logout")
async def logout(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id in sessions:
        username = sessions[session_id]
        if username in user_data:
            del user_data[username]
        del sessions[session_id]
    response = RedirectResponse(url="/")
    response.delete_cookie("session_id")
    return response

# ========================
# APPLICATION PRINCIPALE 
# ========================

@app.get("/app")
async def dashboard(request: Request):
    username = get_username_from_session(request)
    if not username:
        return RedirectResponse(url="/")
    
    # HTML identique à ton code perfection
    return HTMLResponse("""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SynthAI – pour la synthèse intelligente.</title>
<style>
body {
    font-family: Arial;
    background: #f5f5f5;
    margin: 0;
    padding: 0;
}
header {
    background: #3f51b5;
    color: white;
    text-align: center;
    padding: 1.2rem;
    font-size: 1.6rem;
    display: flex;
    justify-content: space-between;
    align-items: center;
}
.header-content {
    display: flex;
    justify-content: space-between;
    align-items: center;
    width: 100%;
    max-width: 900px;
    margin: 0 auto;
}
.user-info {
    font-size: 1rem;
}
.logout-btn {
    background: #ff4757;
    color: white;
    border: none;
    padding: 0.5rem 1rem;
    border-radius: 5px;
    cursor: pointer;
    text-decoration: none;
}
.container {
    max-width: 900px;
    margin: 2.5rem auto;
    background: white;
    padding: 2.5rem;
    border-radius: 12px;
    box-shadow: 0 6px 15px rgba(0,0,0,0.1);
}
.drop-zone {
    border: 2px dashed #3f51b5;
    border-radius: 10px;
    padding: 2.5rem;
    text-align: center;
    cursor: pointer;
    color: #3f51b5;
    margin-bottom: 1.5rem;
    transition: 0.3s;
}
.drop-zone.dragover {
    background: #e8eaf6;
}
.progress-bar {
    width: 0;
    height: 20px;
    background: #3f51b5;
    border-radius: 10px;
    margin-top: 1rem;
    margin-bottom: 2rem;
    transition: width 0.3s;
}
section {
    margin-top: 2.5rem;
}
section h3 {
    font-weight: bold;
    margin-bottom: 0.8rem;
}
section pre, section p {
    background: #e3f2fd;
    padding: 1.5rem;
    border-radius: 8px;
    line-height: 1.6;
    margin-bottom: 1.5rem;
    overflow-x: auto;
}
button {
    background: #3f51b5;
    color: white;
    border: none;
    padding: 0.8rem 1.5rem;
    border-radius: 5px;
    cursor: pointer;
    margin-top: 1rem;
}
button:hover {
    background: #303f9f;
}
#chart-container img {
    max-width: 100%;
    margin-top: 1rem;
    border-radius: 8px;
}
input#question {
    margin-top: 0.5rem;
    width: 80%;
    padding: 0.5rem;
}
</style>
</head>
<body>
<header>
    <div class="header-content">
        <h1>RAG + Analyse Document IA</h1>
        <div class="user-info">
            Connecté en tant que <strong>""" + username + """</strong>
            <a href="/logout" class="logout-btn">Déconnexion</a>
        </div>
    </div>
</header>
<div class="container">
<div class="drop-zone" id="drop-zone">
Glisser-déposer un fichier ou cliquer
<input type="file" id="file-input" style="display:none">
</div>
<div class="progress-bar" id="progress-bar"></div>
<section>
<h3>Aperçu</h3>
<pre id="preview-text"></pre>
</section>
<section>
<h3>Résumé</h3>
<p id="summary-text"></p>
</section>
<section>
<h3>Suggestions</h3>
<p id="suggestions-text"></p>
</section>
<section id="chart-container">
<h3>Graphique</h3>
<img id="chart-img" src="">
</section>
<section>
<button id="download-btn" style="display:none;">Télécharger le fichier corrigé</button>
</section>
<section>
<h3>Poser une question</h3>
<input id="question" style="width:80%;padding:0.5rem;">
<button id="ask-btn">Envoyer</button>
<p id="answer-text"></p>
</section>
</div>
<script>
const dz=document.getElementById("drop-zone");
const fi=document.getElementById("file-input");
const pb=document.getElementById("progress-bar");
const preview=document.getElementById("preview-text");
const summary=document.getElementById("summary-text");
const sugg=document.getElementById("suggestions-text");
const chart=document.getElementById("chart-img");
const btn=document.getElementById("download-btn");
dz.addEventListener("click",()=>fi.click());
dz.addEventListener("dragover",e=>{e.preventDefault();dz.classList.add("dragover");});
dz.addEventListener("dragleave",e=>dz.classList.remove("dragover"));
dz.addEventListener("drop",e=>{e.preventDefault();dz.classList.remove("dragover");handle(e.dataTransfer.files);});
fi.addEventListener("change",e=>handle(e.target.files));
document.addEventListener("dragover",e=>e.preventDefault());
document.addEventListener("drop",e=>e.preventDefault());
function handle(files){if(files.length>0)upload(files[0]);}
function upload(f){
const fd=new FormData();fd.append("file",f);
const xhr=new XMLHttpRequest();xhr.open("POST","/upload",true);
xhr.upload.onprogress=e=>{if(e.lengthComputable)pb.style.width=(e.loaded/e.total*100)+"%";}
xhr.onload=()=>{
    if(xhr.status===200){
const d=JSON.parse(xhr.responseText);
    preview.textContent=d.preview;
    summary.textContent=d.summary;
    sugg.textContent=d.suggestions;
    if(d.chart_base64)chart.src="data:image/png;base64,"+d.chart_base64;
    if(d.corrected_file){btn.style.display="inline-block";btn.onclick=()=>window.open("/download_corrected/"+d.corrected_file,"_blank");}
}else alert("Erreur "+xhr.status);
};
xhr.send(fd);
}
document.getElementById("ask-btn").addEventListener("click",async()=>{
const q=document.getElementById("question").value.trim();
if(!q)return;
const fd=new FormData();fd.append("question",q);
const r=await fetch("/ask",{method:"POST",body:fd});
const d=await r.json();document.getElementById("answer-text").textContent=d.answer;
});
</script>
</body>
</html>""")

# ========================
# ENDPOINTS BACKEND (identique à ton code perfection)
# ========================

@app.post("/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    username = get_username_from_session(request)
    if not username:
        return JSONResponse({"error": "Non authentifié"}, status_code=401)
    
    user_session = user_data[username]
    
    try:
        ext = os.path.splitext(file.filename)[1].lower()
        uid = str(uuid.uuid4())
        file_path = os.path.join(UPLOAD_DIR, f"{username}_{uid}{ext}")
        with open(file_path, "wb") as f:
            f.write(await file.read())
        text = extract_text_from_file(file_path, ext)
        user_session["last_uploaded_text"] = text

        chart_base64 = None
        preview_text = ""
        if ext in [".csv", ".xls", ".xlsx"]:
            user_session["last_uploaded_df"] = pd.read_excel(file_path) if ext != ".csv" else pd.read_csv(file_path)
            preview_text = user_session["last_uploaded_df"].head(10).to_string()
            chart_base64 = dataframe_chart_base64(user_session["last_uploaded_df"])
            corr = correct_excel(user_session["last_uploaded_df"], os.path.join(UPLOAD_DIR, f"{username}_{uid}_corr.xlsx"))
            user_session["last_corrected_file"] = os.path.basename(corr)
        else:
            user_session["last_uploaded_df"] = None
            preview_text = text[:2000]
            corr = correct_docx(file_path, os.path.join(UPLOAD_DIR, f"{username}_{uid}_corr.docx"))
            user_session["last_corrected_file"] = os.path.basename(corr)

        return JSONResponse({
            "preview": preview_text,
            "summary": summarize_with_gemini(text),
            "suggestions": detect_excel_errors(user_session["last_uploaded_df"]) if user_session["last_uploaded_df"] is not None else "",
            "chart_base64": chart_base64,
            "corrected_file": user_session["last_corrected_file"]
        })
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/download_corrected/{filename}")
async def download_corrected(filename: str, request: Request):
    username = get_username_from_session(request)
    if not username:
        return JSONResponse({"error": "Non authentifié"}, status_code=401)
    
    # Vérifier que le fichier appartient à l'utilisateur
    if not filename.startswith(username + "_"):
        return JSONResponse({"error": "Accès non autorisé"}, status_code=403)
    
    path = os.path.join(UPLOAD_DIR, filename)
    return FileResponse(path, filename=filename)

@app.post("/ask")
async def ask_question(request: Request, question: str = Form(...)):
    username = get_username_from_session(request)
    if not username:
        return JSONResponse({"error": "Non authentifié"}, status_code=401)
    
    user_session = user_data.get(username, {})
    if not user_session.get("last_uploaded_text"):
        return JSONResponse({"answer": "Aucun document n'a encore été envoyé."})
    
    answer = gemini_chat([{"role": "user", "content": f"{user_session['last_uploaded_text']}\nRéponds à cette question : {question}"}])
    return JSONResponse({"answer": answer})

# ========================
# LANCEMENT
# ========================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))  # Railway définit le port via cette variable
    uvicorn.run(app, host="0.0.0.0", port=port)
